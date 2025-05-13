
import cv2
import numpy as np
import pytesseract
import os
from docx import Document
from docx.shared import Pt

# Укажите путь к tesseract.exe
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_and_save_to_docx(image_path, output_docx="output.docx", lang='rus'):
    try:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Файл {image_path} не найден")

        img = cv2.imread(image_path)
        if img is None:
            raise ValueError("Не удалось загрузить изображение, проверьте формат файла.")

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.medianBlur(gray, 3)
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        text = pytesseract.image_to_string(gray, lang=lang)
        if not text.strip():
            raise ValueError("Не удалось распознать текст на изображении")

        doc = Document()

        doc.add_heading(f"Результат распознавания: {os.path.basename(image_path)}", level=1)

        # Добавление текста ,размер шрифта 12pt
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(12) 

        doc.save(output_docx)

        print(f"Документ создан: {output_docx}")
        print("\nРаспознанный текст:")
        print(text)
        return text.strip()

    except Exception as e:
        print(f"Ошибка при обработке: {e}")
        return None


if __name__ == "__main__":
    image_path = 'путь к фото.jpg'  # Укажите ваш путь
    output_docx = "распознанный_текст.docx" 

    print(f"файл обрабатывается: {image_path}")
    result_text = extract_and_save_to_docx(image_path, output_docx)

    if not result_text:
        print("\nНе удалось распознать текст")
