import cv2
import numpy as np
import pytesseract
import os
import requests
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt

# Укажите путь к tesseract.exe
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


def download_dejavu_font():
    """Скачивает и устанавливает шрифт DejaVu Sans (оставлено для совместимости)"""
    # Для docx шрифты не нужны, но функция оставлена, если потребуется
    return None


def extract_and_save_to_docx(image_path, output_docx="output.docx", lang='rus'):
    """Функция для распознавания текста и сохранения в DOCX"""
    try:
        # Проверка существования файла
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Файл {image_path} не найден!")

        # Чтение изображения
        img = cv2.imread(image_path)
        if img is None:
            raise ValueError("Не удалось загрузить изображение. Проверьте формат файла.")

        # Улучшение качества изображения
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.medianBlur(gray, 3)
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        # Распознавание текста
        text = pytesseract.image_to_string(gray, lang=lang)
        if not text.strip():
            raise ValueError("Не удалось распознать текст на изображении")

        # Создаем документ Word
        doc = Document()

        # Добавляем заголовок
        doc.add_heading(f"Результат распознавания: {os.path.basename(image_path)}", level=1)

        # Добавляем распознанный текст
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(12)  # Устанавливаем размер шрифта 12pt

        # Сохраняем документ
        doc.save(output_docx)

        print(f"Документ Word успешно создан: {output_docx}")
        print("\nРаспознанный текст:")
        print(text)
        return text.strip()

    except Exception as e:
        print(f"Ошибка при обработке: {e}")
        return None


if __name__ == "__main__":
    image_path = 'photo_5208459711639316327_y.jpg'  # Укажите ваш путь
    output_docx = "recognized_text.docx"  # Имя выходного DOCX файла

    print(f"Обрабатываем файл: {image_path}")
    result_text = extract_and_save_to_docx(image_path, output_docx)

    if not result_text:
        print("\nНе удалось распознать текст")
